#!/usr/bin/env python3
"""
Markdown to DOCX converter
Converts Markdown files with LaTeX math to Word documents
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import sys

def convert_md_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX"""

    # Read Markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new Document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'  # Korean font
    font.size = Pt(11)

    # Split content by lines
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if line.strip() in ['---', '***', '___']:
            i += 1
            continue

        # Headers
        if line.startswith('##'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            # Remove emoji markers like 1️⃣, 2️⃣ etc
            text = re.sub(r'[0-9]\ufe0f?\u20e3', '', text).strip()

            p = doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        # Block quotes
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            p = doc.add_paragraph(text, style='Quote')
            i += 1
            continue

        # Math blocks (display mode)
        if line.strip() == '[' or line.strip() == '\\[':
            # Collect math content
            math_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() not in [']', '\\]']:
                math_lines.append(lines[i])
                i += 1

            if math_lines:
                math_text = '\n'.join(math_lines)
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(f"[수식: {math_text}]")
                run.font.italic = True
            i += 1
            continue

        # Bullet lists
        if line.strip().startswith('*') or line.strip().startswith('-'):
            text = re.sub(r'^\s*[\*\-]\s*', '', line)
            # Handle inline math
            text = re.sub(r'\$([^\$]+)\$', r'[\1]', text)
            # Handle bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\s*\d+\.', line):
            text = re.sub(r'^\s*\d+\.\s*', '', line)
            # Handle inline math
            text = re.sub(r'\$([^\$]+)\$', r'[\1]', text)
            # Handle bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Tables (simple detection)
        if '|' in line and i + 1 < len(lines) and '|' in lines[i+1]:
            # Collect table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                # Parse table
                rows = []
                for tline in table_lines:
                    if re.match(r'^\s*\|[\s\-:]+\|', tline):  # Skip separator
                        continue
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    rows.append(cells)

                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clean math
                            cell_text = re.sub(r'\$([^\$]+)\$', r'[\1]', cell_text)
                            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                            cell.text = cell_text

                    doc.add_paragraph()  # Add space after table
            continue

        # Regular paragraphs
        if line.strip():
            text = line
            # Handle inline math $...$
            text = re.sub(r'\$([^\$]+)\$', r'[\1]', text)
            # Handle bold **...**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Handle italic
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            # Handle code
            text = re.sub(r'`(.+?)`', r'[\1]', text)

            p = doc.add_paragraph(text)
        else:
            # Empty line - add small space
            if i > 0 and i < len(lines) - 1:
                doc.add_paragraph()

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Successfully converted to DOCX file.")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]

    convert_md_to_docx(md_file, docx_file)
